import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT


def export_docx(title: str, sections: list) -> bytes:
    """
    sections = [{"heading": "Title", "content": "text or list"}]
    """
    doc = Document()
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    for section in sections:
        doc.add_heading(section.get("heading", ""), level=2)
        content = section.get("content", "")
        if isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(str(item), style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
        else:
            doc.add_paragraph(str(content))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_pdf(title: str, sections: list) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=inch * 0.75, leftMargin=inch * 0.75,
                            topMargin=inch * 0.75, bottomMargin=inch * 0.75)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('Title', parent=styles['Title'],
                                  fontSize=20, textColor=colors.HexColor('#1a1a2e'),
                                  spaceAfter=6, alignment=TA_CENTER)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'],
                               fontSize=13, textColor=colors.HexColor('#16213e'),
                               spaceBefore=12, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                 fontSize=10, leading=14, spaceAfter=4)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'],
                                   fontSize=10, leading=14, leftIndent=20,
                                   bulletIndent=10, spaceAfter=2)

    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", body_style))
    story.append(Spacer(1, 12))

    for section in sections:
        story.append(Paragraph(section.get("heading", ""), h2_style))
        content = section.get("content", "")
        if isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"• {item}", bullet_style))
        else:
            story.append(Paragraph(str(content), body_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()
